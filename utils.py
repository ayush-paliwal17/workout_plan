import os
import json
import tempfile
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sqlite3
from datetime import *
import tkinter as tk
from PIL import Image,ImageTk
from abc import ABC,abstractmethod

connnection = sqlite3.connect('gym_data.db') #connect to a local database, create a new one if it dosen't exist

cursor = connnection.cursor() #cursor allows us to execute sql commands, it is an interface between the database and sql.

class member_data:

    def create_table():
        cursor.execute("CREATE TABLE IF NOT EXISTS member_data(id_num PRIMARY KEY,Name,Cardio,Batch,Day,Gender,DOB,Address,Phone_Number,e_mail,Fathers Name,Membership Renewal Date)")
        connnection.commit()
        print('Table Created')


    def add_member(member_details):
        #order to be followed (id_num,name,cardio,batch,day,gender,DOB,address,phone number,e_mail,father's name,membership renewal date)
        cursor.execute(f"""
                       INSERT INTO member_data VALUES {member_details}
        """)
        connnection.commit()


    def get_member_details(id_num):
        cursor.execute(f"""
            SELECT * from member_data
            where id_num = '{id_num}'
        """)
        result = cursor.fetchone()
        if result:
            member = {
                "id_num" : id_num,
                "Name" : result[1],
                "Cardio" : result[2],
                "Batch" : result[3],
                "Day" : result[4],
                "Gender" : result[5],
                "DOB" : result[6],
                "Address" : result[7],
                "Phone Number" : result[8],
                "e_mail" : result[9],
                "father's name" : result[10],
                "member renewal date" : result[11],
                "Age" : str(int((date.today()-datetime.strptime(result[6],'%Y-%m-%d').date()).days/365.25))
                    }
            return member
        else:
            raise(Exception)


    def get_all_members():
        cursor.execute(f"""
            SELECT * from member_data
        """)
        result = cursor.fetchall()

        return result

    
    def update_column(key,value,id_value):
        cursor.execute(f"""
            UPDATE member_data
            SET {key} = '{value}'
            WHERE id_num = '{id_value}'
        """)
        connnection.commit()


    def get_last_key():
        cursor.execute(f"""
            SELECT id_num from member_data
        """)
        result = cursor.fetchall()

        return result


class Workout:
    def get_weight_training(ex_id):
        cursor.execute(f"""
            SELECT * from weight_training
            WHERE ex_id = '{ex_id}'
        """)
        result = cursor.fetchone()
        weight_training = [x for x in result[1:]]
        return weight_training

    def get_cardio(ex_id):
        cursor.execute(f"""
            SELECT * from cardio
            where ex_id = {ex_id}
        """)
        result = cursor.fetchone()
        cardio = [x for x in result[1:]]
        return cardio


class Features:
    def print_doc(result_dict,name,part):

        #Initialize The Document
        doc = docx.Document()

        #Page Heading
        # head = doc.add_heading(level = 0).add_run("IRON FITNESS GYM")
        head = doc.add_heading(level = 0)
        head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        head = head.add_run("IRON FITNESS GYM")
        font = head.font
        font.size = Pt(40)

        #Greeting
        para = doc.add_paragraph().add_run(f"Hi {name},\nToday we will be training {part}")
        font = para.font
        font.size = Pt(20)

        #Weight Training
        wt_head = doc.add_heading(level =1).add_run("Weight Training :")
        font = wt_head.font
        font.size = Pt(20)

        for item in result_dict['Weight_Training']:
            para = doc.add_paragraph(style = 'List Number').add_run(item)
            font = para.font
            font.size = Pt(18)

        #Cardio
        cardio = result_dict.get('Cardio')
        if cardio:
            cardio_head = doc.add_heading(level = 1).add_run("Cardio : ")
            font = cardio_head.font
            font.size = Pt(20)    

            for item in cardio:
                para = doc.add_paragraph(style='List Number').add_run(item)
                font = para.font
                font.size = Pt(18)

        #Open Temp file
        filename = tempfile.mktemp(".doc")
        doc.save(filename)
        os.startfile(filename)


    def add_workout_to_DB(id_num,workout):

        #get date
        t_date = str(date.today())

        #read workout history from the json file
        with open("workout_history.json") as f:
            fi_data = json.load(f)

        if fi_data.get(id_num):
            queue = fi_data.get(id_num)
        else:
            queue = []

        queue.append({t_date : workout})
        if len(queue)>30: #limiting the history to 30 unique entries
            queue.pop(0)

        #writing to the json file
        fi_data[id_num] = queue        
        with open("workout_history.json",'w') as f:
            json.dump(fi_data,f)


    def get_history(id_num,dt):
        with open("workout_history.json") as f:
            fi_data = json.load(f)
        queue = fi_data[id_num]
        workout = None
        for history in queue:
            if history.get(dt):
                workout = history.get(dt)
        return workout
    

    def delete_widget(widget_list):
        for widget in widget_list:
            widget.place_forget()
    

    def format_part_list(part_list):
        new_list = list(set(part_list)) #remove duplicates
        new_list.insert(-1,'and')
        if len(new_list) == 2:
            res = new_list[-1]
        else:
            fi = ', '.join(new_list[:-2])
            res = fi+' '+' '.join(new_list[-2:])
        
        return res


class tkinter_templates:
    def home(root):
        #Home Image
        path = r'gallery\home.png'
        img = ImageTk.PhotoImage(Image.open(path))
        panel = tk.Label(root,image=img)
        panel.image = img
        panel.place(relheight=1,relwidth=1)

        #Gym Label
        gym_label = tk.Label(root,text='IRON  FITNESS  GYM',font=('Papyrus',18),background='Black',fg='White')
        gym_label.place(x=(int(700-gym_label.winfo_reqwidth())/2))

        #gym moto label
        moto_label = tk.Label(root,text='Break Your Limits!',font=('Papyrus',12),background='Black',fg='White')
        moto_label.place(x=(int(700-moto_label.winfo_reqwidth())/2),y=45)

    def display(root):
        #display Image
        path = r'gallery\display.png'
        img = ImageTk.PhotoImage(Image.open(path))
        panel = tk.Label(root,image=img)
        panel.image = img
        panel.place(relheight=1,relwidth=1)

        #Gym Label
        gym_label = tk.Label(root,text='IRON  FITNESS  GYM',font=('Papyrus',18),background='Black',fg='White')
        gym_label.place(x=(int(700-gym_label.winfo_reqwidth())/2))

        #gym moto label
        moto_label = tk.Label(root,text='Break Your Limits!',font=('Papyrus',12),background='Black',fg='White')
        moto_label.place(x=(int(700-moto_label.winfo_reqwidth())/2),y=45)


class exercise_list_factory:
    class body_part(ABC):
        @abstractmethod
        def get_exercise_list():
            pass

    class chest(body_part):
        @staticmethod
        def get_exercise_list():
            cursor.execute("Select Exercise Name from Chest")
            res = cursor.fetchall()
            res = [x[0] for x in res]
            return res

    class back(body_part):
        @staticmethod
        def get_exercise_list():
            cursor.execute("Select Exercise Name from Back")
            res = cursor.fetchall()
            res = [x[0] for x in res]
            return res

    class shoulder(body_part):
        @staticmethod
        def get_exercise_list():
            cursor.execute("Select Exercise Name from Shoulders")
            res = cursor.fetchall()
            res = [x[0] for x in res]
            return res

    class biceps(body_part):
        @staticmethod
        def get_exercise_list():
            cursor.execute("Select Exercise Name from Biceps")
            res = cursor.fetchall()
            res = [x[0] for x in res]
            return res

    class triceps(body_part):
        @staticmethod
        def get_exercise_list():
            cursor.execute("Select Exercise Name from Triceps")
            res = cursor.fetchall()
            res = [x[0] for x in res]
            return res

    class legs(body_part):
        @staticmethod
        def get_exercise_list():
            cursor.execute("Select Exercise Name from Legs")
            res = cursor.fetchall()
            res = [x[0] for x in res]
            return res

    class cardio(body_part):
        @staticmethod
        def get_exercise_list():
            cursor.execute("Select Exercise Name from Cardio_list")
            res = cursor.fetchall()
            res = [x[0] for x in res]
            return res

    def part_factory(part):
        part_dict = {
            "Chest" : exercise_list_factory.chest(),
            "Back" : exercise_list_factory.back(),
            "Shoulders" : exercise_list_factory.shoulder(),
            "Biceps" : exercise_list_factory.biceps(),
            "Triceps" : exercise_list_factory.triceps(),
            "Legs" : exercise_list_factory.legs(),
            "Cardio" : exercise_list_factory.cardio()
            }
        return part_dict.get(part)

    def get_exercise_list(part):
        try:
            part_obj = exercise_list_factory.part_factory(part)
            return part_obj.get_exercise_list()
        except Exception:
            raise Exception("Unable to Create Object")

